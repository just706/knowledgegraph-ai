"""文件解析：从上传文件中提取纯文本。

支持：
- TXT / Markdown（直接读取）
- PDF（pypdf 提取文字层；若文字层为空则尝试 OCR 兜底）
- Word（.docx，python-docx 提取段落与表格）
- 图片（png/jpg/jpeg/bmp/gif，pytesseract OCR）

OCR 需要本地安装 Tesseract 引擎，未安装时自动降级：解析函数抛出明确异常，
由调用方返回友好提示，不影响其他类型上传（AI 宪法：优雅降级）。
"""
import io

from pypdf import PdfReader

# 支持的类型 → 标准化后缀
SUPPORTED_TYPES = {
    "txt", "md", "markdown", "pdf", "docx",
    "png", "jpg", "jpeg", "bmp", "gif",
}

# 需要 OCR 的图片类型
IMAGE_TYPES = {"png", "jpg", "jpeg", "bmp", "gif"}

# OCR 是否可用（导入时探测 Tesseract，不可用则降级）
OCR_AVAILABLE = False
try:
    import pytesseract  # noqa: F401
    from PIL import Image  # noqa: F401

    # 探测 Tesseract 引擎是否安装
    pytesseract.get_tesseract_version()
    OCR_AVAILABLE = True
except Exception:  # noqa: BLE001 任意异常都视为 OCR 不可用
    OCR_AVAILABLE = False


def normalize_file_type(filename: str) -> str:
    """从文件名提取标准化类型；不支持时返回空串。"""
    if "." not in filename:
        return ""
    ext = filename.rsplit(".", 1)[-1].lower()
    if ext == "markdown":
        ext = "md"
    return ext if ext in SUPPORTED_TYPES else ""


def extract_text(file_bytes: bytes, file_type: str) -> str:
    """从文件字节提取纯文本。

    Args:
        file_bytes: 文件原始字节
        file_type: 标准化类型（txt/md/pdf/docx/png/jpg/...）
    Returns:
        提取后的纯文本（可能为空串，由调用方处理）
    """
    if file_type in ("txt", "md"):
        return _read_text(file_bytes)

    if file_type == "pdf":
        return _extract_pdf(file_bytes)

    if file_type == "docx":
        return _extract_docx(file_bytes)

    if file_type in IMAGE_TYPES:
        return _extract_image(file_bytes)

    raise ValueError(f"不支持的文件类型: {file_type}")


def _read_text(file_bytes: bytes) -> str:
    try:
        return file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        return file_bytes.decode("utf-8", errors="ignore")


def _extract_pdf(file_bytes: bytes) -> str:
    reader = PdfReader(io.BytesIO(file_bytes))
    parts: list[str] = []
    for page in reader.pages:
        text = page.extract_text() or ""
        parts.append(text)

    text = "\n".join(parts).strip()
    # 文字层为空（扫描件）时，用 OCR 兜底
    if not text and OCR_AVAILABLE:
        try:
            from PIL import Image

            images = _pdf_pages_to_images(file_bytes)
            return "\n".join(_ocr_image(img) for img in images).strip()
        except Exception:  # noqa: BLE001 OCR 失败则回退到空文本
            return ""
    return text


def _pdf_pages_to_images(file_bytes: bytes) -> list:
    """将 PDF 每页渲染为 PIL Image（依赖 pdf2image + poppler，缺失则抛异常）。"""
    from pdf2image import convert_from_bytes

    return convert_from_bytes(file_bytes)


def _extract_docx(file_bytes: bytes) -> str:
    from docx import Document as DocxDocument

    doc = DocxDocument(io.BytesIO(file_bytes))
    parts: list[str] = []

    # 段落
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    # 表格
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def _extract_image(file_bytes: bytes) -> str:
    if not OCR_AVAILABLE:
        raise ValueError(
            "当前环境未安装 Tesseract OCR 引擎，无法识别图片文字。"
            "请安装 Tesseract 后重试，或使用 TXT/PDF/Word 等文本类资料。"
        )
    from PIL import Image

    image = Image.open(io.BytesIO(file_bytes))
    return _ocr_image(image).strip()


def _ocr_image(image) -> str:
    import pytesseract

    # 优先尝试中文+英文；失败回退默认语言
    try:
        return pytesseract.image_to_string(image, lang="chi_sim+eng")
    except Exception:  # noqa: BLE001 指定语言包缺失时回退
        return pytesseract.image_to_string(image)
